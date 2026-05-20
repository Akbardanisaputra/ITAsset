import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()
    
    # Set styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Path to screenshots
    screenshot_dir = 'report_screenshots'
    md_file = 'LAPORAN_PROTOTYPE_UIUX.md'

    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Header 1
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header 2
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        # Header 3
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Bullet points
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Numbered lists
        elif line[0:2].isdigit() and line[2:4] == '. ':
            doc.add_paragraph(line[3:], style='List Number')
        # Bold text (simple detection)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        # Images: ![Caption](path)
        elif line.startswith('![') and '](' in line:
            caption_start = line.find('[') + 1
            caption_end = line.find(']')
            path_start = line.find('(') + 1
            path_end = line.find(')')
            
            caption = line[caption_start:caption_end]
            img_path = line[path_start:path_end]
            
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add caption below image
                p_caption = doc.add_paragraph(caption)
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_caption.style.font.italic = True
                p_caption.style.font.size = Pt(9)
            else:
                doc.add_paragraph(f"[Image Missing: {img_path}]")
        # Captions for images (starts with *Caption:)
        elif line.startswith('*Caption:'):
            p = doc.add_paragraph(line[1:-1]) # Remove * and *
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style.font.italic = True
            p.style.font.size = Pt(9)
        # Horizontal rule
        elif line == '---':
            doc.add_page_break()
        # Normal text
        else:
            doc.add_paragraph(line)

    output_file = 'LAPORAN_PROTOTYPE_UIUX.docx'
    doc.save(output_file)
    print(f"Successfully generated {output_file}")

if __name__ == '__main__':
    create_docx()
